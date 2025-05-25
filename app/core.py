import os
import sys
import win32com.client
import pythoncom
from pyhwpx import Hwp
from docx import Document
from docx.shared import RGBColor
import pyautogui
import time
import win32gui
import win32con

def extract_red_text(input_file):
    doc = Document(input_file)
    all_results = []
    for table in doc.tables:
        table_results = []
        for row in table.rows:
            red_words = []
            for cell in row.cells:
                current_red_word = ""
                in_red = False
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if hasattr(run.font.color, 'rgb') and run.font.color.rgb == RGBColor(255, 0, 0):
                            current_red_word += run.text
                            in_red = True
                        else:
                            if in_red and current_red_word.strip():
                                red_words.append(current_red_word.strip())
                                current_red_word = ""
                                in_red = False
                if in_red and current_red_word.strip():
                    red_words.append(current_red_word.strip())
            if red_words:
                table_results.append(red_words)
        if table_results:
            all_results.append(table_results)
    return all_results

def underline_all_fixed_spaces(doc, num_spaces=12):
    target_text = " " * num_spaces
    rng = doc.Content
    rng.Find.ClearFormatting()
    rng.Find.Text = target_text
    while rng.Find.Execute():
        rng.Font.Underline = True
        rng.Collapse(0)
def create_synonym_questions_in_range(range_obj, red_word_groups):
    """
    주어진 범위에 동의어 문제를 생성하는 함수
    
    Args:
        range_obj: Word 문서의 Range 객체
        red_word_groups: 이중 리스트 형태의 빨간색 단어 그룹
            예: [["단어1", "동의어1", "동의어2"], ["단어2", "동의어3", "동의어4"]]
    """
    first_words = {}
    rest_map = {}
    for group in red_word_groups:
        if not group:
            continue
        first, *rest = group
        first_words[first] = -1
        if rest:
            if isinstance(rest[0], str) and ',' in rest[0]:
                rest = [word.strip() for word in rest[0].split(',')]
            rest_map[first] = rest
    
    word_positions = []
    for word in first_words:
        search_range = range_obj.Duplicate
        search_range.Find.ClearFormatting()
        search_range.Find.Text = word
        search_range.Find.MatchWholeWord = True
        search_range.Find.Forward = True
        search_range.Find.Wrap = 1
        
        if search_range.Find.Execute():
            search_range.Font.Bold = True
            search_range.Font.Italic = True
            search_range.Font.Underline = True
            word_positions.append((word, search_range.Start))
    
    sorted_words = [word for word, _ in sorted(word_positions, key=lambda x: x[1])]
    end_position = range_obj.End
    range_obj.Document.Range(end_position, end_position).InsertAfter("\n")
    end_position += 1 
    
    for word in sorted_words:
        word_range = range_obj.Document.Range(end_position, end_position)
        word_range.InsertAfter(word)
        word_range.Font.Bold = True
        word_range.Font.Italic = True
        word_range.Font.Underline = True
        word_range.Font.Name = "함초롬바탕"
        word_range.Font.Size = 10
        word_range.ParagraphFormat.LineSpacing = 24
        end_position = word_range.End
        question_range = range_obj.Document.Range(end_position, end_position)
        question_range.InsertAfter(" 의 동의어를 쓰시오.")
        question_range.Font.Name = "함초롬바탕"
        question_range.Font.Italic = False
        question_range.Font.Bold = False
        question_range.Font.Underline = False
        question_range.Font.Size = 10
        question_range.ParagraphFormat.LineSpacing = 24
        end_position = question_range.End
        if word in rest_map:
            synonyms = ", ".join(rest_map[word])
            footnote_range = range_obj.Document.Range(end_position, end_position)
            footnote = range_obj.Document.Footnotes.Add(Range=footnote_range, Text=synonyms)
            footnote.Range.Font.Size = 9
            footnote.Range.Font.Name = "맑은 고딕"
            end_position = footnote_range.End + 1  
        range_obj.Document.Range(end_position, end_position).InsertAfter("\n")
        end_position += 1  
        if word in rest_map:
            processed_hints = []
            for hint in rest_map[word]:
                if not hint.strip():
                    continue
                if ' ' in hint:
                    first_letters = [word[0].lower() for word in hint.split()]
                    processed_hint = '    '.join(first_letters) + ' ' * 3
                else:
                    processed_hint = hint[0].lower() + ' ' * 12
                
                processed_hints.append(processed_hint)
            
            hint_text = " → " + ", ".join(processed_hints)
            
            hint_range = range_obj.Document.Range(end_position, end_position)
            hint_range.InsertAfter(hint_text)
            hint_range.Font.Name = "함초롬바탕"
            hint_range.Font.Size = 10
            hint_range.ParagraphFormat.LineSpacing = 24
            end_position = hint_range.End
            range_obj.Document.Range(end_position, end_position).InsertAfter("\n")
            end_position += 1
        range_obj.Document.Range(end_position, end_position).InsertAfter("\n")
        end_position += 1 

def create_synonym_questions_from_red_text(word_file_path, red_word_groups):
    word_app = None
    doc = None
    
    try:
        pythoncom.CoInitialize()
        
        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
        except Exception as e:
            print(f"Word 애플리케이션 초기화 실패: {str(e)}")
            import subprocess
            subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
        
        doc = word_app.Documents.Open(
            word_file_path,
            ReadOnly=False,
            Visible=False,
            NoEncodingDialog=True
        )
        tables = doc.Tables

        if tables.Count > 0:
            for i in range(tables.Count):
                current_table = tables.Item(i + 1)
                
                if i == 0:
                    content_between_tables = doc.Range(0, current_table.Range.Start - 1)
                else:
                    previous_table = tables.Item(i)
                    content_between_tables = doc.Range(previous_table.Range.End + 1, current_table.Range.Start - 1)
                
                create_synonym_questions_in_range(content_between_tables, red_word_groups[i])
        
        for table in tables:
            table.Delete()

        underline_all_fixed_spaces(doc)
        underline_all_fixed_spaces(doc, 7)
        underline_all_fixed_spaces(doc, 3)
        
        output_file_path = os.path.splitext(word_file_path)[0] + "_synonym_questions.docx"
        doc.SaveAs(output_file_path)
        
        return output_file_path
        
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        raise
        
    finally:
        try:
            if doc:
                doc.Content.Copy()
                doc.Close(SaveChanges=False)
            if word_app:
                word_app.Quit()
        except Exception as e:
            print(f"리소스 정리 중 오류: {str(e)}")
        finally:
            pythoncom.CoUninitialize()


def save_as_hwp(input_file_path):
    try:    
        output_file_path = os.path.splitext(input_file_path)[0] + ".hwp"
        
        if os.path.exists(output_file_path):
            os.remove(output_file_path)
        
        hwp = Hwp()        
        pyautogui.hotkey('ctrl', 'alt', 'v')
        time.sleep(0.5)
        pyautogui.press(['down', 'down', 'enter'])
        time.sleep(0.5)
        
        hwp.save_as(output_file_path)
        hwp.quit()
        time.sleep(0.5)

        if not os.path.exists(output_file_path):
            raise Exception("HWP 파일 저장 실패")
            
        return output_file_path
        
    except Exception as e:
        print(f"HWP 변환 중 오류 발생: {str(e)}")
        raise

def hwp_to_docx(input_file_path: str) -> str:
    if not input_file_path.lower().endswith(('.hwp', '.hwpx')):
        raise ValueError("입력 파일은 .hwp 또는 .hwpx 형식이어야 합니다.")
        
    hwp = Hwp()
    hwp.open(input_file_path)
    time.sleep(0.5)
    pyautogui.hotkey('ctrl', 'a')
    pyautogui.hotkey('ctrl', 'c')
    hwp.clear()
    hwp.quit()

    file_name = os.path.basename(input_file_path) 
    output_file_path = os.path.join(os.getcwd(), file_name.replace('.hwp', '.docx').replace('.hwpx', '.docx'))  # 현재 경로에 저장
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = True
    
    doc = word.Documents.Add()
    time.sleep(1)
    
    def callback(hwnd, extra):
        if win32gui.IsWindowVisible(hwnd):
            title = win32gui.GetWindowText(hwnd)
            if "Word" in title:
                win32gui.SetForegroundWindow(hwnd)
                win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                return False
        return True
    
    win32gui.EnumWindows(callback, None)
    time.sleep(0.5)
    
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(2)
    pyautogui.press('ctrl')
    time.sleep(1)
    pyautogui.press('k')
    time.sleep(0.5)
    
    doc.SaveAs(output_file_path)
    doc.Close()
    word.Quit()

    return output_file_path
